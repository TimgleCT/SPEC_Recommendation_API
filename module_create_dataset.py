import os
import docx
import docx2txt
import re
import jieba
import numpy as np
import pandas as pd
from tqdm import tqdm
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import TfidfTransformer
from sklearn.decomposition import PCA
from sklearn.cluster import AgglomerativeClustering
import matplotlib.pyplot as plt
import seaborn as sns;sns.set()

jieba.case_sensitive = True
jieba.set_dictionary('jieba_dict/dict.txt.big.txt')
jieba.load_userdict('jieba_dict/tableName.txt')
jieba.load_userdict('jieba_dict/programNameList.txt')
jieba.load_userdict('jieba_dict/allvocb.txt')

tableNameFile = open('jieba_dict/tableName.txt', 'r', encoding='UTF-8')
tableNameList = tableNameFile.read().split('\n')
programNameFile = open('jieba_dict/programNameList.txt', 'r', encoding='UTF-8')
programNameList = programNameFile.read().split('\n')


def getRelatedFilePar(doc, dic, dicKeyName, TargetList):

    findChinese = re.compile("[^A-Za-z0-9^.]")
    
    findRelatedFile = False
    content = []
    for para in doc.paragraphs:
      contentList = " ".join(jieba.cut(para.text)).split(" ")
      for i in range(len(contentList)):
        if contentList[i] in TargetList and contentList[i] not in content:
          content.append(contentList[i])
    
    dic[dicKeyName] = ",".join(content)

    return dic



def getRelatedFile(doc, dic, dicKeyName, tableCol, TargetList):
    
    fileList = []
    findChinese = re.compile("[^A-Za-z0-9^.^_]")
    
    for table in doc.tables:  # 遍歷所有表格
        fileNameIndex = 0
        for cell in table.rows[0].cells:
            if cell.text in tableCol:
                for file in table.rows[1:]:
                    try:
                      content = file.cells[fileNameIndex].text
                      if len(findChinese.findall(content)) == 0:
                        content = "".join(findChinese.sub('',file.cells[fileNameIndex].text))
                        if content != "" and content in TargetList:
                          fileList.append(content)
                    except:
                      fileList = []
                      print("表格格式讀取錯誤")
                    
                dic[dicKeyName] = ",".join(fileList)
            
            fileNameIndex += 1

    if len(fileList) > 0:
        return dic
    else:
        if dicKeyName == "相關檔案":
            return getRelatedFilePar(doc, dic, "相關檔案", TargetList)
        elif dicKeyName == "相關模組":
            return getRelatedFilePar(doc, dic, "相關模組", TargetList)
        else:
            print("錯誤!")



def getTableModel():
    
    path = 'dataset/spec/'
    docTableModel = []

    for file in tqdm(os.listdir(path)):
        if file.endswith(".docx"): 
            
            filename = path + file
            Doc = docx.Document(filename)
            TableModel = {}
            TableModel["文件檔名"] = file
            TableModel = getRelatedFile(Doc, TableModel, "相關檔案", ["檔案名稱","中文說明","檔案內容"], tableNameList)
            TableModel = getRelatedFile(Doc, TableModel, "相關模組", ["程式名稱","CLASS"], programNameList)
                
            docTableModel.append(TableModel)
            
    return docTableModel



def getTableModelTFIDF(docDictList):

    allDocTable = []
    allDocModel = []
    fileIndex = []

    for i in range(len(docDictList)):
        tableStr = docDictList[i]["相關檔案"]
        modelStr = docDictList[i]["相關模組"]
        fileIndexStr = docDictList[i]["文件檔名"]
        fileIndex.append(fileIndexStr)
        allDocTable.append(tableStr)
        allDocModel.append(modelStr)

    vectorizer = CountVectorizer()

    T = vectorizer.fit_transform(allDocTable).toarray()
    T_word = vectorizer.get_feature_names()

    M = vectorizer.fit_transform(allDocModel).toarray()
    M_word = vectorizer.get_feature_names()

    print(T)
    print(len(T_word))
    print(T_word)

    print(M)
    print(len(M_word))
    print(M_word)

    return T, M, fileIndex


def getRelationMatrix(cosThreshold, TFIDFMatrix, fileIndex, csvName):

    # 剛剛計算的資料表TFIDF矩陣
    TFIDFMatrix_1 = TFIDFMatrix
    # 因自己與自己有關係，要扣掉
    relateCount = -1 * len(fileIndex)
    relateMatrix = []

    for vect_1 in TFIDFMatrix_1:
        reshapeVect = np.array(vect_1).reshape(1, -1)
        cosScoreMatrix = cosine_similarity(TFIDFMatrix, reshapeVect)
        cosScoreMatrix = cosScoreMatrix.reshape(cosScoreMatrix.shape[0])

        # 計算後的cos相似度矩陣為vect_1與其他vect的cos相似度。若相似度大於cosThreshold則視為相關
        for i in range(len(cosScoreMatrix)):
            if cosScoreMatrix[i] >= cosThreshold:
                cosScoreMatrix[i] = 1
                relateCount += 1
            else:
                cosScoreMatrix[i] = 0

        relateMatrix.append(cosScoreMatrix)

    # 以Dataframe儲存成關聯矩陣並轉成csv檔
    relateMatrix_df = pd.DataFrame(relateMatrix,columns=fileIndex,index=fileIndex)
    relateMatrix_df.to_csv(csvName)
    print(relateMatrix_df.head())
    print("有多少個文件彼此關聯(不含自己)：",relateCount/2)

    return relateMatrix


def preprocessSPEC(source_directory):

    findEnNum = re.compile("[^A-Za-z0-9]")
    rule = re.compile("[^\u4e00-\u9fa5]")

    docList = []
    docIndex = []
    stopwordcount = 0
    getstopword = []

    stopWordFile = open('jieba_dict/stop_word.txt', 'r', encoding='UTF-8')
    stopWordList = stopWordFile.read().split('\n') 

    for file in tqdm(os.listdir(source_directory)):
        noStopWordTokenList = []
        if file.endswith(".docx"):
            filename = source_directory + file
            docIndex.append(file)
            docxToText = docx2txt.process(filename)

            docxToTextEnNum = "".join(findEnNum.sub('',docxToText))
            docxToText = "".join(rule.sub('',docxToText))

            tokenizeStr = " ".join(jieba.cut(docxToText))
            tokenizeStrEnNum = " ".join(jieba.cut(docxToTextEnNum))

            newTokenStr = tokenizeStr.split(" ")
            newTokenStrEnNum = tokenizeStrEnNum.split(" ")


            for token in newTokenStr:
                if token in stopWordList:
                    stopwordcount += 1
                    getstopword.append(token)
                else:
                    noStopWordTokenList.append(token)

            for token in newTokenStrEnNum:
                if token in tableNameList or token in programNameList:
                    noStopWordTokenList.append(token)

            newTokenizeStr = " ".join(noStopWordTokenList)
            docList.append(newTokenizeStr)

    return docList



def docxTFIDF(docList):
    vectorizer = CountVectorizer()
    X = vectorizer.fit_transform(docList)
    word = vectorizer.get_feature_names()
    print(len(word))
    transformer = TfidfTransformer()
    tfidf = transformer.fit_transform(X)
    weight = tfidf.toarray()
    print(weight)

    return weight


def AgglomerativeCluster(weight, PCA_dimension, threshold):

    pca = PCA(n_components = PCA_dimension)
    tfPCA = pca.fit_transform(weight)

    pcaVisual = PCA(n_components = 2)
    visualXY = pcaVisual.fit_transform(weight)

    clustering = AgglomerativeClustering(n_clusters=None, distance_threshold=threshold).fit(tfPCA)

    fig = plt.figure(figsize=(12,10))
    ax1 = fig.add_subplot(2, 2, 1)
    ax1.set_title("y_AgglomerativeClustering_PCA_First")
    ax1.scatter(visualXY[:, 0], visualXY[:, 1], c=clustering.labels_, cmap='jet')

    plt.show()

    aggNumCluster = len(set(clustering.labels_))
    aggClusterResult = clustering.labels_
    cosScore = getClusterCosScore(getClusterDocVec(weight, aggNumCluster, aggClusterResult))
    print("總共群數：", aggNumCluster)
    print("本次分群平均得分：", cosScore)

    return aggClusterResult, round(cosScore, 2)


def getClusterDocVec(weight, num_cluster, output):

    resultDic = {}
    for i in range(num_cluster):
        resultDic[str(i)] = []

    for index, class_ in enumerate(output):
        resultDic[str(class_)].append(weight[index].tolist())
  
    return resultDic


def getClusterCosScore(ClusterDocVecDic):

    totalCosScore = 0
    # 群迴圈
    for i in range(len(ClusterDocVecDic)):
        # 群內迴圈
        cosScore = 0
        flag = 0
        loopCount = 0
        for j in tqdm(range(len(ClusterDocVecDic[str(i)]))):
            for q in range(flag, len(ClusterDocVecDic[str(i)])):
                if j == q:
                    continue
                else:
                    reshapeVect_j = np.array(ClusterDocVecDic[str(i)][j]).reshape(1, -1)
                    reshapeVect_q = np.array(ClusterDocVecDic[str(i)][q]).reshape(1, -1)
                    cosScoreMatrix = cosine_similarity(reshapeVect_j, reshapeVect_q)
                    cosScoreMatrix = cosScoreMatrix.reshape(cosScoreMatrix.shape[0])
                    cosScore += cosScoreMatrix[0]
                    loopCount += 1
            flag += 1
        try:
            cosScore /= loopCount
            print("本群分數：",cosScore)
        except:
            cosScore = cosScore

        totalCosScore += cosScore

    return totalCosScore/len(ClusterDocVecDic) 

def findBestThreshold(minThreshold, maxThreshold, minScore, fileIndex):
    pca = PCA(n_components = 256)
    tfPCA = pca.fit_transform(weight)
    bestScore = 0
    bestClusterNum = 0
    bestThreshold = 0
    for num in range(minThreshold, maxThreshold):
        threshold = num/10
        clustering = AgglomerativeClustering(n_clusters=None,distance_threshold=threshold).fit(tfPCA)
        aggCluster = clustering.labels_
        clusterNum = len(set(aggCluster))
        score = getClusterCosScore(getClusterDocVec(clusterNum, aggCluster, fileIndex))
        if threshold > bestThreshold and  score > minScore:
            bestClusterNum = clusterNum
            bestThreshold = threshold
            bestScore = score
        print("threshold為"+ str(threshold) +"時，cluster數為"+ str(clusterNum) +"時，群內相似分數為"+ str(score))
        print("---------------------------------")

    print("最佳threshold：",bestThreshold)
    print("cluster數：",bestClusterNum)
    print("相似分數：",bestScore)

    return bestThreshold



def getRelateMatrixFromCluster(clusterResult, fileIndex, csvName):

    # clusterResult為分群結果
    # csvName為設定輸出關聯矩陣的csv檔名

    # 因自己與自己有關係，要扣掉
    relateCount = -1 * len(fileIndex)
    clusterRelateMatrix = []
    for i in clusterResult:
        relateList = []
        for j in clusterResult:

            # 如果i跟j都屬於同一群的話則視為兩者有相關
            if i == j:
                relateList.append(1)
                relateCount += 1
            else:
                relateList.append(0)

        # 將與i的相關矩陣存入clusterRelateMatrix中
        clusterRelateMatrix.append(relateList)

    # 以Dataframe儲存成關聯矩陣並轉成csv檔
    clusterRelate_df = pd.DataFrame(clusterRelateMatrix,index=fileIndex,columns=fileIndex)
    clusterRelate_df.to_csv(csvName)
    print(clusterRelate_df)
    print("有多少個文件彼此關聯(不含自己)：",relateCount/2)

    # 回傳clusterRelateMatrix
    return clusterRelateMatrix



def outputFinalRelateMatrix(clusterRelateMatrix, relateMatrix_model, relateMatrix_table, fileIndex):

    # clusterRelateMatrix, relateMatrix_model, relateMatrix_table為分群關聯矩陣、模組關聯矩陣、資料表關聯矩陣

    # 因自己與自己有關係，要扣掉
    relateCount = -1 * len(fileIndex)

    # 宣告新關聯矩陣
    relateMatrix_result = []

    for i in range(len(fileIndex)):

        # 宣告某SPEC與其他SPEC的關聯陣列：relateMatrix_row
        relateMatrix_row = []
        for j in range(len(fileIndex)):

            # 若分群關聯矩陣、模組關聯矩陣、資料表關聯矩陣內的其中一個元素為1則視為相關
            if clusterRelateMatrix[i][j] == 1 or relateMatrix_model[i][j] == 1 or relateMatrix_table[i][j] == 1:
                relateMatrix_row.append(1)
                relateCount += 1
            else:
                relateMatrix_row.append(0)
    
        # 將結果存入新關聯矩陣中
        relateMatrix_result.append(relateMatrix_row)

    # 以Dataframe儲存成關聯矩陣並轉成csv檔
    resultRelate_df = pd.DataFrame(relateMatrix_result,index=fileIndex,columns=fileIndex)
    resultRelate_df.to_csv('new_dataset/resultRelate_df.csv')
    print(resultRelate_df)
    print("有多少個文件彼此關聯(不含自己)：",relateCount/2)

    return relateMatrix_result



def createDataset(relateMatrix_result, fileIndex, csvName):

    # relateMatrix_result最後的關聯矩陣
    # csvName 輸出的資料集名稱設定

    # 宣告資料集陣列
    dataset = []

    # flag 若過去已有紀錄的關聯便不再重複紀錄EX：A與B相關，則遇到B與A相關時便不再紀錄
    flag = 0
    for i in range(len(relateMatrix_result)):
        for j in range(flag, len(relateMatrix_result[0])):

            # 自己與自己的關聯不記錄
            if i == j:
                continue
            else:
                dataset.append([fileIndex[i],fileIndex[j],relateMatrix_result[i][j]])
        flag += 1

    # 以Dataframe儲存成資料集並轉成csv檔
    dataset_df = pd.DataFrame(dataset,columns=["SPEC1","SPEC2","relation"])
    print(dataset_df.head())
    dataset_df.to_csv(csvName)